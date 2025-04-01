#!/usr/bin/env python3
"""
Document Processor with Google Vision AI and Tesseract fallback
Handles PDF, images, screenshots, and Word documents
"""
import os
import io
import mimetypes
from typing import Optional, Dict, Any, List, Union
from google.cloud import vision
import pytesseract
from PIL import Image
import cv2
import numpy as np
import PyPDF2
import pdf2image
import docx
from dotenv import load_dotenv

# Load environment variables
load_dotenv()


class DocumentProcessor:
    def __init__(self, vision_credentials_path: Optional[str] = None):
        """Initialize the document processor.

        Args:
            vision_credentials_path: Path to Google Cloud Vision credentials JSON file.
                                  If None, will use GOOGLE_APPLICATION_CREDENTIALS from env.
        """
        self.vision_client = None
        self.fallback_to_tesseract = (
            os.getenv("FALLBACK_TO_TESSERACT", "true").lower() == "true"
        )

        try:
            if vision_credentials_path:
                os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = vision_credentials_path
            self.vision_client = vision.ImageAnnotatorClient()
            print("Google Cloud Vision initialized successfully")
        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Failed to initialize Google Cloud Vision: {e}")
                print("Falling back to Tesseract OCR")
            else:
                raise

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process any supported document type and extract text.

        Args:
            file_path: Path to the document file.

        Returns:
            Dictionary containing extracted text and metadata.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        mime_type, _ = mimetypes.guess_type(file_path)

        try:
            if mime_type == "application/pdf":
                return self._process_pdf(file_path)
            elif mime_type in [
                "image/jpeg",
                "image/png",
                "image/tiff",
                "image/bmp",
                "image/gif",
            ]:
                return self._process_image(file_path)
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return self._process_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Primary processing failed: {e}")
                print("Attempting fallback processing...")
                return self._fallback_process(file_path)
            else:
                raise

    def _process_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Process PDF file using Google Cloud Vision."""
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(pdf_path)

            # Process each page
            extracted_text = []
            page_metadata = []

            for i, image in enumerate(images):
                # Save image temporarily
                temp_path = f"temp_page_{i}.png"
                image.save(temp_path)

                try:
                    # Process with Vision API
                    result = self._process_image_with_vision(temp_path)
                    extracted_text.append(result["text"])
                    page_metadata.append(result["metadata"])
                finally:
                    # Clean up
                    if os.path.exists(temp_path):
                        os.remove(temp_path)

            return {
                "text": "\n\n".join(extracted_text),
                "metadata": {
                    "type": "pdf",
                    "pages": len(images),
                    "page_metadata": page_metadata,
                },
            }

        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: PDF processing with Vision failed: {e}")
                print("Falling back to PyPDF2")
                with open(pdf_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    text = "\n".join(page.extract_text() for page in reader.pages)
                return {
                    "text": text,
                    "metadata": {
                        "type": "pdf",
                        "pages": len(reader.pages),
                        "processing_method": "fallback",
                    },
                }
            else:
                raise

    def _process_image(self, image_path: str) -> Dict[str, Any]:
        """Process image file using Google Cloud Vision."""
        try:
            # Open and preprocess image
            with Image.open(image_path) as img:
                # Convert to RGB if needed
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")

                # Save to temporary buffer
                with io.BytesIO() as buffer:
                    img.save(buffer, format="PNG")
                    content = buffer.getvalue()

            return self._process_image_with_vision(content)

        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Image processing with Vision failed: {e}")
                print("Falling back to Tesseract")
                return self._process_image_with_tesseract(image_path)
            else:
                raise

    def _process_docx(self, docx_path: str) -> Dict[str, Any]:
        """Process Word document."""
        try:
            doc = docx.Document(docx_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

            return {
                "text": text,
                "metadata": {"type": "docx", "paragraphs": len(doc.paragraphs)},
            }
        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Word document processing failed: {e}")
                print("Falling back to image-based OCR")
                # Convert Word to PDF then to image
                # TODO: Implement Word to PDF conversion if needed
                raise
            else:
                raise

    def _process_image_with_vision(
        self, image_content: Union[str, bytes]
    ) -> Dict[str, Any]:
        """Process image content using Google Cloud Vision API."""
        if self.vision_client is None:
            raise Exception("Google Cloud Vision client not initialized")

        # Create Vision API image object
        if isinstance(image_content, str):
            with open(image_content, "rb") as image_file:
                content = image_file.read()
        else:
            content = image_content

        image = vision.Image(content=content)

        # Perform text detection
        response = self.vision_client.text_detection(image=image)

        if response.error.message:
            raise Exception(
                "{}\nFor more info on error messages, check: "
                "https://cloud.google.com/apis/design/errors".format(
                    response.error.message
                )
            )

        # Extract text and metadata
        texts = response.text_annotations
        if not texts:
            return {"text": "", "metadata": {"confidence": 0.0}}

        # Get confidence scores
        confidences = [
            text.confidence for text in texts[1:]
        ]  # Skip first text (full text)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        return {
            "text": texts[0].description,
            "metadata": {
                "confidence": round(avg_confidence, 3),
                "text_blocks": len(texts) - 1,  # Subtract 1 for the full text
            },
        }

    def _process_image_with_tesseract(self, image_path: str) -> Dict[str, Any]:
        """Process image using Tesseract OCR."""
        try:
            # Use OpenCV to improve image quality
            img = cv2.imread(image_path)
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            img = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

            # Use pytesseract
            text = pytesseract.image_to_string(img)

            return {
                "text": text,
                "metadata": {"processing_method": "tesseract", "image_size": img.shape},
            }
        except Exception as e:
            raise Exception(f"Tesseract OCR failed: {str(e)}")

    def _fallback_process(self, file_path: str) -> Dict[str, Any]:
        """Fallback processing method when primary methods fail."""
        try:
            # Try to process as image first
            return self._process_image_with_tesseract(file_path)
        except Exception as e:
            print(f"Fallback processing failed: {e}")
            raise

    def detect_document_properties(self, file_path: str) -> Dict[str, Any]:
        """Detect additional document properties using Google Cloud Vision API."""
        try:
            if self.vision_client is None:
                raise Exception("Google Cloud Vision client not initialized")

            mime_type, _ = mimetypes.guess_type(file_path)

            if mime_type == "application/pdf":
                # Convert first page to image
                images = pdf2image.convert_from_path(
                    file_path, first_page=1, last_page=1
                )
                if not images:
                    raise Exception("Failed to convert PDF to image")

                # Save first page temporarily
                temp_path = "temp_first_page.png"
                images[0].save(temp_path)
                try:
                    return self._detect_properties_from_image(temp_path)
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
            elif mime_type in [
                "image/jpeg",
                "image/png",
                "image/tiff",
                "image/bmp",
                "image/gif",
            ]:
                return self._detect_properties_from_image(file_path)
            else:
                raise ValueError(
                    f"Unsupported file type for property detection: {mime_type}"
                )

        except Exception as e:
            return {"error": str(e), "processing_method": "fallback"}

    def _detect_properties_from_image(self, image_path: str) -> Dict[str, Any]:
        """Detect properties from an image file."""
        with io.open(image_path, "rb") as image_file:
            content = image_file.read()

        image = vision.Image(content=content)

        # Detect text and document properties
        doc_response = self.vision_client.document_text_detection(image=image)

        # Get confidence scores and language
        confidence = 0.0
        language = "unknown"
        pages = 0

        if doc_response.full_text_annotation:
            # Get overall confidence
            if doc_response.full_text_annotation.pages:
                confidences = [
                    page.confidence for page in doc_response.full_text_annotation.pages
                ]
                confidence = sum(confidences) / len(confidences) if confidences else 0.0
                pages = len(doc_response.full_text_annotation.pages)

            # Get detected language
            for page in doc_response.full_text_annotation.pages:
                if page.property and page.property.detected_languages:
                    language = page.property.detected_languages[0].language_code
                    break

        return {
            "language": language,
            "confidence": round(confidence, 3),
            "pages": pages,
            "error": (
                None if not doc_response.error.message else doc_response.error.message
            ),
        }


# Example usage
if __name__ == "__main__":
    # Initialize processor
    processor = DocumentProcessor()

    # Test files
    test_files = [
        "sample_term_sheet.pdf",
        "sample_term_sheet.png",  # You'll need to create this
        "sample_term_sheet.docx",  # You'll need to create this
    ]

    # Process each file
    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"\nProcessing {file_path}...")
            try:
                # Process document
                result = processor.process_document(file_path)
                print("\nExtracted text (first 500 chars):")
                print("-" * 50)
                print(
                    result["text"][:500] + "..."
                    if len(result["text"]) > 500
                    else result["text"]
                )

                print("\nMetadata:")
                print("-" * 50)
                print(result["metadata"])

                # Detect properties
                properties = processor.detect_document_properties(file_path)
                print("\nDocument properties:")
                print("-" * 50)
                print(properties)

            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
        else:
            print(f"\nSkipping {file_path} (file not found)")
