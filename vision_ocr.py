import os
from typing import Optional, Dict, Any, List
from google.cloud import vision
from dotenv import load_dotenv
import io
from PIL import Image
import pdf2image
import docx
import tempfile
import mimetypes

# Load environment variables
load_dotenv()


class VisionOCR:
    def __init__(self, credentials_path: Optional[str] = None):
        """Initialize the Google Cloud Vision client.

        Args:
            credentials_path: Optional path to credentials JSON file.
                            If not provided, will use GOOGLE_APPLICATION_CREDENTIALS from env.
        """
        if credentials_path:
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = credentials_path

        try:
            self.client = vision.ImageAnnotatorClient()
            self.fallback_to_tesseract = (
                os.getenv("FALLBACK_TO_TESSERACT", "true").lower() == "true"
            )
        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Failed to initialize Google Cloud Vision: {e}")
                print("Falling back to Tesseract OCR")
                self.client = None
            else:
                raise

    def extract_text(self, file_path: str) -> str:
        """Extract text from any supported document type.

        Args:
            file_path: Path to the document file.

        Returns:
            Extracted text as string.
        """
        mime_type, _ = mimetypes.guess_type(file_path)

        if mime_type == "application/pdf":
            return self.extract_text_from_pdf(file_path)
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return self.extract_text_from_docx(file_path)
        elif mime_type in [
            "image/jpeg",
            "image/png",
            "image/tiff",
            "image/bmp",
            "image/gif",
        ]:
            return self.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from a Word document.

        Args:
            docx_path: Path to the Word document.

        Returns:
            Extracted text as string.
        """
        try:
            doc = docx.Document(docx_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Word document processing failed: {e}")
                print("Falling back to image-based OCR")
                # Convert Word to PDF then to image
                with tempfile.NamedTemporaryFile(
                    suffix=".pdf", delete=False
                ) as temp_pdf:
                    # TODO: Add Word to PDF conversion if needed
                    # For now, we'll just raise the error
                    raise Exception("Word to PDF conversion not implemented yet")
            else:
                raise

    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from an image using Google Cloud Vision API.

        Args:
            image_path: Path to the image file.

        Returns:
            Extracted text as string.
        """
        try:
            if self.client is None:
                raise Exception("Google Cloud Vision client not initialized")

            # Open and preprocess image
            with Image.open(image_path) as img:
                # Convert to RGB if needed
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")

                # Save to temporary buffer
                with io.BytesIO() as buffer:
                    img.save(buffer, format="PNG")
                    content = buffer.getvalue()

            image = vision.Image(content=content)
            response = self.client.text_detection(image=image)

            if response.error.message:
                raise Exception(
                    "{}\nFor more info on error messages, check: "
                    "https://cloud.google.com/apis/design/errors".format(
                        response.error.message
                    )
                )

            texts = response.text_annotations
            if not texts:
                return ""

            return texts[0].description

        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: Google Cloud Vision failed: {e}")
                print("Falling back to Tesseract OCR")
                # Import tesseract here to avoid dependency if not needed
                import pytesseract

                return pytesseract.image_to_string(Image.open(image_path))
            else:
                raise

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from a PDF using Google Cloud Vision API.

        Args:
            pdf_path: Path to the PDF file.

        Returns:
            Extracted text as string.
        """
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(pdf_path)

            # Extract text from each image
            extracted_text = []
            for i, image in enumerate(images):
                # Save image temporarily
                temp_path = f"temp_page_{i}.png"
                image.save(temp_path)

                # Extract text
                text = self.extract_text_from_image(temp_path)
                extracted_text.append(text)

                # Clean up
                os.remove(temp_path)

            return "\n\n".join(extracted_text)

        except Exception as e:
            if self.fallback_to_tesseract:
                print(f"Warning: PDF processing with Google Cloud Vision failed: {e}")
                print("Falling back to PyPDF2")
                import PyPDF2

                with open(pdf_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    return "\n".join(page.extract_text() for page in reader.pages)
            else:
                raise

    def detect_document_properties(self, file_path: str) -> Dict[str, Any]:
        """Detect additional document properties using Google Cloud Vision API.

        Args:
            file_path: Path to the file (PDF, image, or Word document).

        Returns:
            Dictionary containing document properties.
        """
        try:
            if self.client is None:
                raise Exception("Google Cloud Vision client not initialized")

            mime_type, _ = mimetypes.guess_type(file_path)

            if mime_type == "application/pdf":
                # Convert first page to image
                images = pdf2image.convert_from_path(
                    file_path, first_page=1, last_page=1
                )
                if not images:
                    raise Exception("Failed to convert PDF to image")

                # Save first page temporarily
                temp_path = "temp_first_page.png"
                images[0].save(temp_path)
                try:
                    return self._detect_document_properties_from_image(temp_path)
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
            elif mime_type in [
                "image/jpeg",
                "image/png",
                "image/tiff",
                "image/bmp",
                "image/gif",
            ]:
                return self._detect_document_properties_from_image(file_path)
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                # For Word documents, we'll return basic properties
                doc = docx.Document(file_path)
                return {
                    "language": "unknown",
                    "confidence": 1.0,
                    "pages": len(doc.sections),
                    "error": None,
                }
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")

        except Exception as e:
            return {
                "language": "unknown",
                "confidence": 0.0,
                "pages": 0,
                "error": str(e),
            }

    def _detect_document_properties_from_image(self, image_path: str) -> Dict[str, Any]:
        """Internal method to detect properties from an image file.

        Args:
            image_path: Path to the image file.

        Returns:
            Dictionary containing document properties.
        """
        with io.open(image_path, "rb") as image_file:
            content = image_file.read()

        image = vision.Image(content=content)

        # Detect text and document properties
        doc_response = self.client.document_text_detection(image=image)

        # Get confidence scores and language
        confidence = 0.0
        language = "unknown"
        pages = 0

        if doc_response.full_text_annotation:
            # Get overall confidence
            if doc_response.full_text_annotation.pages:
                confidences = [
                    page.confidence for page in doc_response.full_text_annotation.pages
                ]
                confidence = sum(confidences) / len(confidences) if confidences else 0.0
                pages = len(doc_response.full_text_annotation.pages)

            # Get detected language
            for page in doc_response.full_text_annotation.pages:
                if page.property and page.property.detected_languages:
                    language = page.property.detected_languages[0].language_code
                    break

        return {
            "language": language,
            "confidence": round(confidence, 3),
            "pages": pages,
            "error": (
                None if not doc_response.error.message else doc_response.error.message
            ),
        }
